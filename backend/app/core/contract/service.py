"""合同体检主流程（M7.5）：上传合同 → 报告。

串联：解析(parser) → 抽取(extractor) → 合规比对(checker) → 处置建议(remedy)
→ 体检报告 JSON。

报告结构：
{
  "report_id": "...",
  "file_name": "...",
  "summary": {"total": N, "violations": n, "warnings": n, "ok": n, "not_specified": n},
  "findings": [{field, text, verdict, basis, reason, risk}],
  "remedies": {remedies: [...], complex_advice: "..."},
  "disclaimer": "本报告仅供参考，不构成法律意见……"
}
"""
from io import BytesIO
from typing import Any

from app.core.contract.check import check_all
from app.core.contract.extractor import extract_clauses
from app.core.contract.remedy import generate_remedies
from datetime import datetime

DISCLAIMER = (
    "本体检报告由 AI 自动生成，仅供参考，不构成法律意见。"
    "涉及具体权益纠纷，请咨询专业律师或当地法律援助机构。"
    "如需法律文书，请由专业律师起草。"
)

def run_contract_check(contract_text: str, file_name: str = "合同") -> dict[str, Any]:
    """完整体检流程，返回报告 JSON。"""
    # 1. 条款抽取
    clauses = extract_clauses(contract_text)

    # 2. 合规比对 + 风险分级
    findings = check_all(clauses, contract_context=contract_text)

    # 3. 处置建议（维权路径 + 证据清单）
    remedies = generate_remedies(findings)

    # 4. 汇总
    summary = {
        "total": len(findings),
        "violations": sum(1 for f in findings if f["verdict"] == "违法"),
        "warnings": sum(1 for f in findings if f["verdict"] == "模糊"),
        "ok": sum(1 for f in findings if f["verdict"] == "合法"),
        "not_specified": sum(1 for f in findings if f["verdict"] == "未约定"),
    }

    return {
        "report_id": file_name,
        "file_name": file_name,
        "summary": summary,
        "findings": findings,
        "remedies": remedies,
        "disclaimer": DISCLAIMER,
    }




def _build_report_docx(report: dict) -> bytes:
    """把体检报告渲染成 Word 文档,返回 docx 文件 bytes。"""
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    doc = Document()
    # 全局正文样式:宋体 11 磅(中文字体需显式设 eastAsia,否则显示异常)
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 1. 标题 + 生成时间
    doc.add_heading(f"《{report.get('file_name', '合同')}》体检报告", level=0)
    doc.add_paragraph(f"生成时间:{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # 2. 体检结论
    summary = report.get("summary", {})
    doc.add_heading("体检结论", level=1)
    doc.add_paragraph(
        f"共检查 {summary.get('total', 0)} 项条款:违法 {summary.get('violations', 0)} 项、"
        f"模糊 {summary.get('warnings', 0)} 项、合规 {summary.get('ok', 0)} 项、"
        f"未约定 {summary.get('not_specified', 0)} 项"
    )

    # 3. 逐条检查结果
    doc.add_heading("条款明细", level=1)
    for i, f in enumerate(report.get("findings", []), 1):
        p = doc.add_paragraph()
        # 条款名 + 结论(加粗)
        run = p.add_run(f"{i}. {f.get('field', '')} —— {f.get('verdict', '')}")
        run.bold = True
        # 原文 / 依据 / 风险(有则写)
        if f.get("text"):
            doc.add_paragraph(f"条款原文:{f['text']}")
        if f.get("basis"):
            doc.add_paragraph(f"检查依据:{f['basis']}")
        if f.get("reason"):
            doc.add_paragraph(f"风险说明:{f['reason']}")

    # 4. 处置建议(每项分 问题/路径/证据/提醒 四段渲染,避免 JSON 原文)
    remedies = report.get("remedies", {})
    if remedies:
        doc.add_heading("处置建议", level=1)
        for r in remedies.get("remedies", []):
            # 问题简述(加粗)
            p = doc.add_paragraph()
            run = p.add_run(f"问题:{r.get('issue', '')}")
            run.bold = True
            # 维权路径
            if r.get("path"):
                doc.add_paragraph(f"维权路径:{r['path']}")
            # 证据清单(顿号连接)
            if r.get("evidence"):
                doc.add_paragraph(f"证据准备:{'、'.join(r['evidence'])}")
            # 特别提醒
            if r.get("note"):
                doc.add_paragraph(f"提醒:{r['note']}")
        if remedies.get("complex_advice"):
            doc.add_paragraph(f"复杂情况:{remedies['complex_advice']}")

    # 5. 免责声明
    doc.add_paragraph(report.get("disclaimer", DISCLAIMER))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

