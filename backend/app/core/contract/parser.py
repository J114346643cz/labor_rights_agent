from io import BytesIO
from pathlib import Path

SUPPORTED_SUFFIXES = {".md", ".txt", ".markdown", ".pdf", ".docx"}

def _clean(text: str) -> str:
    """清洗：合并多余空行、去除首尾空白。
    有些 md 文件复制粘贴出来，会有一大堆空行，向量化之前要压缩，减少无效 chunk。
    """
    import re
    # 把连续3个及以上换行，替换成2个换行
    text = re.sub(r"\n{3,}","\n\n",text)
    # 去掉字符串最开头、最末尾的空格、换行
    return text.strip()

def _decode_text(content: bytes) -> str:
    """文本文件解码：优先 UTF-8，失败回退 GBK（Windows 常见）。"""
    for enc in ("utf-8", "gbk"):
        try:
            text = content.decode(enc)
            return _clean(text)
        except UnicodeDecodeError:
            continue
    # 兜底：替换非法字符
    return _clean(content.decode("utf-8", errors="replace"))


def _extract_pdf(content: bytes) -> str:
    """PDF 提取：pdfplumber 逐页取文本，拼接。"""
    try:
        import pdfplumber
    except ImportError:
        raise ValueError("PDF 解析依赖未安装，请运行: uv sync")

    pages_text = []
    with pdfplumber.open(BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages_text.append(text)
    full = "\n\n".join(pages_text)
    if not full.strip():
        raise ValueError("PDF 未提取到文本（可能是扫描件/图片型 PDF，暂不支持 OCR）")
    return _clean(full)


def _extract_docx(content: bytes) -> str:
    """Word 提取：python-docx 取段落文本（含表格）。"""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("Word 解析依赖未安装，请运行: uv sync")

    try:
        doc = Document(BytesIO(content))
    except Exception as e:
        # 伪 docx(文本文件改名/网页另存/微信传输)不是 zip 包,解析必然失败
        # 兜底:先按文本解码(UTF-8/GBK),能解出内容就直接用,救回这类文件
        try:
            text = _decode_text(content)
            if text.strip():
                return text
        except Exception:
            pass
        raise ValueError(
            "文件无法解析:既不是有效的 Word 文档(.docx),也不是文本文件。\n"
            "请确认:① Word 文件请保存为 .docx 格式(不支持旧版 .doc) "
            "② 纯文本请用 .txt/.md 扩展名"
        ) from e
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    full = "\n".join(parts)
    if not full.strip():
        raise ValueError("Word 文档未提取到文本")
    return _clean(full)

def extract_text(filename: str, content: bytes) -> str:
    """从上传文件提取纯文本。"""
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"不支持的文件类型 {suffix}，支持：{sorted(SUPPORTED_SUFFIXES)}")

    if suffix in (".md", ".txt", ".markdown"):
        return _decode_text(content)

    if suffix == ".pdf":
        return _extract_pdf(content)

    if suffix == ".docx":
        return _extract_docx(content)

    raise ValueError(f"未处理的文件类型 {suffix}")